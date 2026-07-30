"""
Pipeline Manager — orchestrates the full document analysis flow.
Persists results to the DB and drives the Job Queue.
"""
from __future__ import annotations

import logging
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import AsyncSessionLocal
from app.infrastructure.database.models_intelligence import (
    DocumentAnalysis, AnalysisStatus, LayoutElement as LayoutElementModel,
    ExtractedTable, ExtractedEntity, AISuggestion, ProcessingJob
)
from app.infrastructure.database.models import File
from app.services.classification.document_classifier import DocumentClassifier
from app.services.pipeline.base_pipeline import PipelineContext
from app.services.pipeline.pipelines.generic_pipeline import GenericPipeline
from app.services.pipeline.pipelines.invoice_pipeline import InvoicePipeline
from app.services.pipeline.pipelines.contract_pipeline import ContractPipeline
from app.services.pipeline.pipelines.cv_pipeline import CVPipeline
from app.services.pipeline.pipelines.bank_statement_pipeline import BankStatementPipeline
from app.jobs.job_queue import register_handler

logger = logging.getLogger(__name__)

_classifier = DocumentClassifier()

# Pipeline registry — keyed by doc_type
_PIPELINES: dict = {
    "invoice":        InvoicePipeline(),
    "receipt":        InvoicePipeline(),      # reuse invoice pipeline
    "contract":       ContractPipeline(),
    "resume":         CVPipeline(),
    "bank_statement": BankStatementPipeline(),
}
_generic_pipeline = GenericPipeline()


def _get_pipeline(doc_type: str):
    return _PIPELINES.get(doc_type, _generic_pipeline)


# ── Job handler (registered with the queue) ───────────────────────────────────

@register_handler("analysis")
async def handle_analysis_job(payload: dict) -> dict:
    """Entry point called by the job worker."""
    file_id: int = payload["file_id"]
    file_path: str = payload["file_path"]
    file_format: str = payload.get("file_format", "")
    analysis_id: int = payload["analysis_id"]

    async with AsyncSessionLocal() as db:
        # Mark analysis as running
        analysis = await db.get(DocumentAnalysis, analysis_id)
        if not analysis:
            raise ValueError(f"DocumentAnalysis {analysis_id} not found")

        analysis.status = AnalysisStatus.RUNNING
        analysis.updated_at = datetime.now(timezone.utc)
        await db.commit()

        t0 = time.monotonic()
        try:
            # ── Step 1: Classify ──────────────────────────────────────────
            text_for_classify = _quick_text(file_path, file_format)
            clf_result = _classifier.classify(text_for_classify, Path(file_path).name)

            analysis.doc_type = clf_result.doc_type
            analysis.doc_type_confidence = clf_result.confidence
            analysis.language = clf_result.language
            await db.commit()

            # ── Step 2: Run pipeline ──────────────────────────────────────
            pipeline = _get_pipeline(clf_result.doc_type)
            ctx = PipelineContext(
                file_id=file_id,
                file_path=file_path,
                file_format=file_format,
                analysis_id=analysis_id,
                extra={"doc_type": clf_result.doc_type},
            )
            ctx = pipeline.run(ctx)

            # ── Step 3: Persist results ───────────────────────────────────
            analysis.raw_text = ctx.raw_text[:50000] if ctx.raw_text else None
            analysis.language = ctx.language
            analysis.page_count = ctx.page_count or analysis.page_count
            analysis.has_tables = ctx.has_tables
            analysis.has_images = ctx.has_images
            analysis.pipeline_used = pipeline.name
            analysis.model_versions = {"ocr": "tesseract-5", "layout": "pdfplumber-1", "ner": "regex-1"}

            # Layout elements
            for elem in ctx.layout_elements:
                db.add(LayoutElementModel(
                    analysis_id=analysis_id,
                    page_number=elem.get("page_number", 1),
                    element_type=elem.get("element_type", "other"),
                    x1=elem.get("x1"), y1=elem.get("y1"),
                    x2=elem.get("x2"), y2=elem.get("y2"),
                    confidence=elem.get("confidence"),
                    content=elem.get("content", "")[:500] if elem.get("content") else None,
                    meta=elem.get("meta", {}),
                ))

            # Extracted tables
            for tbl in ctx.tables:
                cells_json = [
                    {"row": c.row, "col": c.col, "value": c.value,
                     "rowspan": c.rowspan, "colspan": c.colspan}
                    for c in tbl.cells
                ]
                db.add(ExtractedTable(
                    analysis_id=analysis_id,
                    page_number=tbl.page_number,
                    row_count=tbl.row_count,
                    col_count=tbl.col_count,
                    has_header=tbl.has_header,
                    has_merged_cells=tbl.has_merged_cells,
                    table_data=cells_json,
                    headers=tbl.headers,
                    confidence=tbl.confidence,
                ))

            # Entities
            for ent in ctx.entities:
                db.add(ExtractedEntity(
                    analysis_id=analysis_id,
                    entity_type=ent.get("entity_type", ""),
                    value=ent.get("value", "")[:500],
                    normalized_value=ent.get("normalized_value", "")[:500] or None,
                    confidence=ent.get("confidence"),
                    context=ent.get("context", "")[:200] or None,
                    page_number=ent.get("page_number"),
                ))

            # Suggestions
            for sug in ctx.suggestions:
                db.add(AISuggestion(
                    file_id=file_id,
                    analysis_id=analysis_id,
                    suggestion_type=sug.get("suggestion_type", ""),
                    title=sug.get("title_ar") or sug.get("title"),
                    description=sug.get("description"),
                    action_params=sug.get("action_params", {}),
                    priority=sug.get("priority", 5),
                ))

            duration_ms = int((time.monotonic() - t0) * 1000)
            analysis.processing_ms = duration_ms
            analysis.status = AnalysisStatus.COMPLETED
            analysis.updated_at = datetime.now(timezone.utc)
            await db.commit()

            # ── Step 4: Auto-index for search ─────────────────────────
            try:
                from app.services.search.search_service import search_service
                file_row = await db.get(File, file_id)
                await search_service.index_document(
                    db,
                    file_id=file_id,
                    analysis_id=analysis_id,
                    user_id=file_row.owner_id if file_row else 0,
                    text=ctx.raw_text or "",
                    doc_type=clf_result.doc_type,
                    language=ctx.language,
                    filename=file_row.original_name if file_row else "",
                )
            except Exception as idx_exc:
                logger.warning("Search indexing failed for file %d: %s", file_id, idx_exc)

            logger.info("Analysis %d completed in %dms (type=%s)", analysis_id, duration_ms, clf_result.doc_type)
            return {
                "analysis_id": analysis_id,
                "doc_type": clf_result.doc_type,
                "language": ctx.language,
                "has_tables": ctx.has_tables,
                "table_count": len(ctx.tables),
                "entity_count": len(ctx.entities),
                "duration_ms": duration_ms,
            }

        except Exception as exc:
            analysis.status = AnalysisStatus.FAILED
            analysis.error_message = str(exc)[:500]
            analysis.updated_at = datetime.now(timezone.utc)
            await db.commit()
            raise


def _quick_text(file_path: str, file_format: str) -> str:
    """Extract a small amount of text quickly for classification."""
    ext = file_format.lower().lstrip(".")
    try:
        if ext == "pdf":
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                pages = pdf.pages[:3]
                return "\n".join(p.extract_text() or "" for p in pages)[:3000]
        elif ext in {"docx", "doc"}:
            import docx
            doc = docx.Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs[:30])[:3000]
        elif ext in {"txt"}:
            return open(file_path, encoding="utf-8", errors="ignore").read(3000)
    except Exception:
        pass
    return ""
